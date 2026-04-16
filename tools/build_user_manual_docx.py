from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "manual_assets"
OUTPUT_DOCX = DOCS_DIR / "Codex-Memory-Board-User-Manual-ZH.docx"

CANVAS_WIDTH = 1600
CANVAS_HEIGHT = 900
BG = "#FBFAF7"
INK = "#1F2937"
MUTED = "#64748B"
LINE = "#CBD5E1"
ACCENT = "#176B87"
ACCENT_2 = "#D97706"
ACCENT_3 = "#0F766E"
BOX_FILL = "#FFFFFF"


@dataclass
class Box:
    x: int
    y: int
    w: int
    h: int
    text: str
    fill: str = BOX_FILL
    outline: str = ACCENT


def load_font(size: int, bold: bool = False):
    font_candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for candidate in font_candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_TITLE = load_font(40, bold=True)
FONT_SUBTITLE = load_font(28, bold=True)
FONT_BODY = load_font(24)
FONT_SMALL = load_font(20)


def ensure_output_dirs() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def new_canvas(title: str, subtitle: str):
    image = Image.new("RGB", (CANVAS_WIDTH, CANVAS_HEIGHT), BG)
    draw = ImageDraw.Draw(image)
    draw.text((80, 48), title, font=FONT_TITLE, fill=INK)
    draw.text((80, 110), subtitle, font=FONT_SMALL, fill=MUTED)
    draw.line((80, 160, CANVAS_WIDTH - 80, 160), fill=LINE, width=3)
    return image, draw


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font,
    fill=INK,
    width=20,
    line_gap=10,
):
    x, y = xy
    for raw_line in text.splitlines():
        parts = wrap(raw_line, width=width, break_long_words=False, replace_whitespace=False) or [""]
        for part in parts:
            draw.text((x, y), part, font=font, fill=fill)
            y += font.size + line_gap
    return y


def draw_box(draw: ImageDraw.ImageDraw, box: Box, font=FONT_BODY, text_fill=INK):
    draw.rounded_rectangle(
        (box.x, box.y, box.x + box.w, box.y + box.h),
        radius=24,
        fill=box.fill,
        outline=box.outline,
        width=4,
    )
    lines: list[str] = []
    for raw_line in box.text.splitlines():
        width_guess = max(8, int(box.w / max(font.size, 18)) + 2)
        wrapped = wrap(raw_line, width=width_guess, break_long_words=False, replace_whitespace=False) or [""]
        lines.extend(wrapped)
    line_height = font.size + 8
    total_height = len(lines) * line_height
    current_y = box.y + max(12, (box.h - total_height) // 2)
    for line in lines:
        left, top, right, bottom = draw.textbbox((0, 0), line, font=font)
        text_width = right - left
        current_x = box.x + (box.w - text_width) // 2
        draw.text((current_x, current_y), line, font=font, fill=text_fill)
        current_y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=ACCENT, width=6):
    draw.line((start, end), fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - 22 * direction, ey - 12), (ex - 22 * direction, ey + 12)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - 12, ey - 22 * direction), (ex + 12, ey - 22 * direction)]
    draw.polygon(points, fill=color)


def build_workflow_diagram() -> Path:
    path = ASSETS_DIR / "workflow_overview_cn.png"
    image, draw = new_canvas("图 1  项目记忆工作闭环", "从初始化到交接，再到自检，帮助长任务稳定续跑")
    boxes = [
        Box(90, 280, 200, 120, "cmb init", fill="#E0F2FE", outline=ACCENT),
        Box(340, 280, 200, 120, "cmb status", fill="#F8FAFC", outline=ACCENT),
        Box(590, 280, 200, 120, "cmb log", fill="#FFF7ED", outline=ACCENT_2),
        Box(840, 280, 200, 120, "cmb next", fill="#ECFDF5", outline=ACCENT_3),
        Box(1090, 280, 200, 120, "cmb handoff", fill="#FEF3C7", outline=ACCENT_2),
        Box(1340, 280, 180, 120, "cmb validate", fill="#F8FAFC", outline=ACCENT),
    ]
    for box in boxes:
        draw_box(draw, box, FONT_SUBTITLE)
    for index in range(len(boxes) - 1):
        draw_arrow(
            draw,
            (boxes[index].x + boxes[index].w, boxes[index].y + boxes[index].h // 2),
            (boxes[index + 1].x, boxes[index + 1].y + boxes[index + 1].h // 2),
        )
    draw_arrow(draw, (1430, 400), (1430, 620))
    draw_arrow(draw, (1430, 620), (440, 620))
    draw_arrow(draw, (440, 620), (440, 400))
    draw_wrapped_text(
        draw,
        (90, 710),
        "阅读方式：先用 status 看当前局面，执行过程中用 log 回写状态。\n"
        "不确定下一步时用 next，准备切换线程时用 handoff。\n"
        "担心文档结构已经失真时，用 validate 做自检。",
        FONT_SMALL,
        fill=MUTED,
        width=40,
        line_gap=12,
    )
    image.save(path)
    return path


def build_architecture_diagram() -> Path:
    path = ASSETS_DIR / "architecture_layers_cn.png"
    image, draw = new_canvas("图 2  项目分层架构", "命令入口、编排、规则、解析、存储和 Markdown 契约各自负责一层")
    layers = [
        Box(220, 185, 1160, 100, "CLI 层\ncli.py / __main__.py\n负责命令注册、参数接收和用户输出", fill="#E0F2FE"),
        Box(220, 310, 1160, 120, "Board 编排层\ninit_memory.py / documentation_board.py / next_board.py\nhandoff_board.py / validate_board.py\n负责把多文件读取、解析和规则调用串起来", fill="#F8FAFC"),
        Box(220, 455, 1160, 100, "规则层\nnext_step.py / handoff.py / validate.py\n负责下一步建议、交接文本和校验规则判断", fill="#ECFDF5", outline=ACCENT_3),
        Box(220, 580, 1160, 100, "解析与模型层\nparser.py / models.py\n负责把 Markdown 文本转成稳定的结构化对象", fill="#FEF3C7", outline=ACCENT_2),
        Box(220, 705, 1160, 100, "存储与文件层\nstore.py + 五个核心 Markdown 文件\n负责文件读写和状态落盘", fill="#FFF7ED", outline=ACCENT_2),
    ]
    for layer in layers:
        draw_box(draw, layer, FONT_SMALL)
    for index in range(len(layers) - 1):
        draw_arrow(
            draw,
            (800, layers[index].y + layers[index].h),
            (800, layers[index + 1].y),
            color=ACCENT,
        )
    image.save(path)
    return path


def build_file_contract_diagram() -> Path:
    path = ASSETS_DIR / "file_contract_map_cn.png"
    image, draw = new_canvas("图 3  五个核心文件各自负责什么", "你可以把这五个文件看成项目状态面板的五个固定区域")
    center = Box(635, 340, 330, 110, "项目状态面板", fill="#E0F2FE", outline=ACCENT)
    draw_box(draw, center, FONT_SUBTITLE)
    nodes = [
        Box(80, 185, 320, 120, "AGENTS.md\n协作规则\n工作边界\n交接注意事项", fill="#F8FAFC"),
        Box(80, 540, 320, 120, "Prompt.md\n提示词片段\n交接模板\n复用表达", fill="#F8FAFC"),
        Box(1200, 185, 320, 120, "Plan.md\n里程碑\n可交付项\n阶段顺序", fill="#ECFDF5", outline=ACCENT_3),
        Box(1200, 540, 320, 120, "Implement.md\n当前开发焦点\n将改哪些文件\n怎么验证", fill="#FFF7ED", outline=ACCENT_2),
        Box(560, 650, 480, 150, "Documentation.md\n当前阶段 / 已完成内容 / 下一步\n最近决策 / 最近验证 / 日志记录", fill="#FEF3C7", outline=ACCENT_2),
    ]
    for node in nodes:
        draw_box(draw, node, FONT_SMALL)
    arrow_pairs = [
        ((400, 245), (635, 360)),
        ((400, 600), (635, 430)),
        ((1200, 245), (965, 360)),
        ((1200, 600), (965, 430)),
        ((800, 650), (800, 450)),
    ]
    for start, end in arrow_pairs:
        draw_arrow(draw, start, end, color=LINE, width=5)
    image.save(path)
    return path


def build_next_rule_diagram() -> Path:
    path = ASSETS_DIR / "next_rule_flow_cn.png"
    image, draw = new_canvas("图 4  cmb next 的判断顺序", "它用固定规则给出建议，所以输出稳定、可解释、便于调试")
    rule_boxes = [
        Box(100, 190, 560, 82, "1. 读取 Documentation.md 和 Plan.md", fill="#E0F2FE", outline=ACCENT),
        Box(100, 300, 560, 82, "2. 如果 Next Step 已写明，就直接输出它", fill="#F8FAFC", outline=ACCENT),
        Box(100, 410, 560, 82, "3. 如果最近一次验证失败，就先修复验证失败", fill="#FFF7ED", outline=ACCENT_2),
        Box(100, 520, 560, 82, "4. 如果当前里程碑还有未完成项，就输出第一条未完成项", fill="#ECFDF5", outline=ACCENT_3),
        Box(100, 630, 560, 82, "5. 如果当前里程碑完成，就看下一里程碑的第一条任务", fill="#F8FAFC", outline=ACCENT),
    ]
    for box in rule_boxes:
        draw_box(draw, box, FONT_SMALL)
    for index in range(len(rule_boxes) - 1):
        draw_arrow(
            draw,
            (rule_boxes[index].x + rule_boxes[index].w // 2, rule_boxes[index].y + rule_boxes[index].h),
            (rule_boxes[index + 1].x + rule_boxes[index + 1].w // 2, rule_boxes[index + 1].y),
            color=ACCENT,
        )
    summary_box = Box(
        820,
        320,
        610,
        260,
        "输出内容保持稳定\n\n当前阶段\n判断依据\n建议的下一步\n\n信息不足时会明确报错",
        fill="#FEF3C7",
        outline=ACCENT_2,
    )
    draw_box(draw, summary_box, FONT_SMALL)
    draw_wrapped_text(
        draw,
        (820, 640),
        "这条命令适合长期维护的原因在于：规则固定，判断过程可解释。\n"
        "你遇到输出不理想时，也能顺着规则顺序快速排查。",
        FONT_SMALL,
        fill=MUTED,
        width=28,
        line_gap=12,
    )
    image.save(path)
    return path


def build_adoption_diagram() -> Path:
    path = ASSETS_DIR / "adoption_flow_cn.png"
    image, draw = new_canvas("图 5  把它用到你自己的项目时怎么落地", "这张图适合第一次接触该工具的使用者")
    steps = [
        Box(120, 220, 300, 120, "1. 安装工具并进入项目目录", fill="#E0F2FE", outline=ACCENT),
        Box(500, 220, 260, 120, "2. 运行\ncmb init", fill="#F8FAFC", outline=ACCENT),
        Box(840, 220, 320, 120, "3. 填写\nAGENTS.md / Plan.md /\nDocumentation.md", fill="#ECFDF5", outline=ACCENT_3),
        Box(260, 450, 360, 130, "4. 用 cmb status 和 cmb next\n确认当前局面、下一步和阶段位置", fill="#FFF7ED", outline=ACCENT_2),
        Box(760, 450, 520, 130, "5. 开发时坚持写 cmb log\n切换线程前运行 cmb handoff\n定期用 cmb validate 检查文档结构", fill="#F8FAFC", outline=ACCENT),
    ]
    for step in steps:
        draw_box(draw, step, FONT_SMALL)
    draw_arrow(draw, (420, 280), (500, 280), color=ACCENT)
    draw_arrow(draw, (760, 280), (840, 280), color=ACCENT)
    draw_arrow(draw, (1000, 340), (440, 450), color=ACCENT)
    draw_arrow(draw, (620, 515), (760, 515), color=ACCENT)
    draw_wrapped_text(
        draw,
        (120, 690),
        "填写文件时，先写清楚三件事：当前阶段、已经完成什么、接下来做什么。\n"
        "这三件事清楚后，status、next、handoff、validate 的效果都会稳定很多。",
        FONT_SMALL,
        fill=MUTED,
        width=45,
        line_gap=12,
    )
    image.save(path)
    return path


def set_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size, bold in [
        ("Title", 20, True),
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11.5, True),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(31, 41, 55)

    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9.5)

    set_page_number(section)


def add_paragraph(document: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = color
    return paragraph


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_code_block(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "CodeBlock"
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F8FAFC")
    cell._tc.get_or_add_tcPr().append(shading)
    document.add_paragraph("")


def add_image(document: Document, image_path: Path, caption: str) -> None:
    document.add_picture(str(image_path), width=Inches(6.6))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    run.italic = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Codex Memory Board\n中文使用说明书")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("面向第一次接触者、项目维护者和准备迁移到自己仓库的使用者")
    run.font.size = Pt(12)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("文档内容包括：用途说明、架构说明、命令说明、文件填写建议、调试方法和中文流程图。")
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph("")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("建议阅读顺序：先看第 1 章到第 4 章，再按需要查看第 5 章到第 10 章。")
    run.italic = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_page_break()


def add_file_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["文件", "主要作用", "建议写什么", "常见问题"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    rows = [
        ("AGENTS.md", "写清团队如何协作", "目标、边界、规则、交接注意事项", "写得太空，会让后续线程不知道哪些事不能做"),
        ("Prompt.md", "存放可复用提示词", "当前任务背景、交接模板、复用短语", "内容过长，下一轮很难快速抓重点"),
        ("Plan.md", "管理阶段和可交付项", "当前里程碑、阶段顺序、每阶段任务清单", "任务过粗或过细，都会影响 next 的建议质量"),
        ("Implement.md", "记录当前开发焦点", "近期修改点、决策、将改哪些文件、怎么验证", "长期不更新，容易和真实代码脱节"),
        ("Documentation.md", "维护当前项目状态", "当前阶段、已完成事项、下一步、最近决策、最近验证、日志", "标题层级被改坏后，status、log、next 会受影响"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph("")


def build_document(diagrams: list[Path]) -> Document:
    document = Document()
    configure_document(document)
    add_title_page(document)

    document.add_heading("1. 这份文档适合谁", level=1)
    add_paragraph(document, "这份说明文档适合三类使用者：第一次接触该项目的人、准备把它迁移到自己仓库的人、以及后续维护命令和文档的人。")
    add_bullets(
        document,
        [
            "如果你只想先学会怎么用，可以重点看第 2 章、第 3 章、第 5 章和第 8 章。",
            "如果你还关心内部代码边界，可以重点看第 4 章、第 6 章和第 7 章。",
            "如果你打算把它套到自己的项目里，可以重点看第 8 章、第 9 章和第 10 章。",
        ],
    )

    document.add_heading("2. 项目能做什么", level=1)
    add_paragraph(document, "Codex Memory Board 是一个本地 Python CLI，用来把项目状态写进固定结构的 Markdown 文件中。")
    add_paragraph(document, "它很适合开发工作分多轮进行、不同线程之间需要快速接手、以及团队希望把项目当前状态写成可读可检查文档的场景。")
    add_bullets(
        document,
        [
            "初始化一套项目记忆文件",
            "读取当前阶段、已完成事项、下一步、最近决策和最近验证结果",
            "把新的开发动作作为结构化日志写回 Documentation.md",
            "根据 Documentation.md 和 Plan.md 给出下一步建议",
            "生成下一轮会话可直接使用的中文交接提示词",
            "检查当前项目记忆是否仍然完整，是否还能支撑后续命令继续工作",
        ],
    )
    add_paragraph(document, "你也可以把它理解成一套轻量的项目状态架构。它既包含命令工具，也包含固定文件契约，还包含一条稳定的工作闭环。")
    add_image(document, diagrams[0], "图 1  项目记忆工作闭环")

    document.add_heading("3. 你在使用时会接触到什么", level=1)
    add_paragraph(document, "日常使用时，主要会接触两层内容：一层是五个核心 Markdown 文件，另一层是六个命令。")
    add_numbered(
        document,
        [
            "先用 cmb init 生成骨架文件。",
            "进入工作前，用 cmb status 了解当前局面。",
            "完成一个清晰动作后，用 cmb log 把变更写回 Documentation.md。",
            "当你需要判断下一步时，用 cmb next 读取计划和状态。",
            "准备切换到下一次会话时，用 cmb handoff 生成交接提示词。",
            "怀疑文档已经失真时，用 cmb validate 做完整性检查。",
        ],
    )

    document.add_heading("4. 当前项目的整体架构", level=1)
    add_paragraph(document, "为了让命令长期保持稳定，这个项目把职责分成几层，每一层只做自己那一段工作。")
    add_image(document, diagrams[1], "图 2  项目分层架构")
    add_bullets(
        document,
        [
            "CLI 层负责命令入口和参数组织。",
            "Board 编排层负责把读取文件、调用解析器、调用规则和整理结果串起来。",
            "规则层负责 next、handoff、validate 里的固定判断逻辑。",
            "解析层负责把 Markdown 结构转换成结构化对象。",
            "存储层只负责读写文本文件。",
        ],
    )

    document.add_heading("5. 六个命令分别怎么用", level=1)
    document.add_heading("5.1 cmb init", level=2)
    add_paragraph(document, "作用：在目标目录下生成五个核心 Markdown 文件。默认不覆盖已存在文件。")
    add_code_block(
        document,
        """
cmb init
cmb init --path path/to/project
        """,
    )
    document.add_heading("5.2 cmb status", level=2)
    add_paragraph(document, "作用：读取 Documentation.md，并把当前阶段、已完成事项、下一步、最近决策和最近验证结果输出出来。")
    add_code_block(
        document,
        """
cmb status
cmb status --path path/to/project
        """,
    )
    document.add_heading("5.3 cmb log", level=2)
    add_paragraph(document, "作用：向 Documentation.md 追加一条结构化日志，并同步刷新最近决策和最近验证两个区域。")
    add_code_block(
        document,
        """
cmb log --item "完成 API 适配" --decision "保留规则驱动方案" --reason "便于长期维护" --verify-command "pytest -q" --verify-result "20 passed"
        """,
    )
    document.add_heading("5.4 cmb next", level=2)
    add_paragraph(document, "作用：根据 Documentation.md 和 Plan.md 的固定规则给出当前最合理的下一步。")
    add_image(document, diagrams[3], "图 4  cmb next 的判断顺序")
    document.add_heading("5.5 cmb handoff", level=2)
    add_paragraph(document, "作用：生成一段中文提示词，供下一轮 Codex 会话直接接手。")
    document.add_heading("5.6 cmb validate", level=2)
    add_paragraph(document, "作用：检查文件是否存在、结构是否完整、状态是否足以支撑后续命令继续工作。")

    document.add_heading("6. 五个核心文件应该写什么", level=1)
    add_paragraph(document, "如果你只想先抓住最重要的点，可以先记住：Plan.md 管计划，Documentation.md 管当前状态，AGENTS.md 管协作边界。")
    add_image(document, diagrams[2], "图 3  五个核心文件各自负责什么")
    add_file_table(document)

    document.add_heading("7. 项目目录和代码模块怎么读", level=1)
    add_paragraph(document, "下面这棵目录树足够帮助第一次读代码的人建立整体印象。")
    add_code_block(
        document,
        """
codex-memory-board/
├─ AGENTS.md
├─ Prompt.md
├─ Plan.md
├─ Implement.md
├─ Documentation.md
├─ pyproject.toml
├─ docs/
│  ├─ PROJECT_GUIDE.md
│  ├─ ARCHITECTURE_DIAGRAMS.md
│  └─ Codex-Memory-Board-中文使用说明书.docx
├─ src/
│  └─ codex_memory_board/
│     ├─ cli.py
│     ├─ parser.py
│     ├─ store.py
│     ├─ models.py
│     ├─ documentation_board.py
│     ├─ next_board.py
│     ├─ next_step.py
│     ├─ handoff_board.py
│     ├─ handoff.py
│     ├─ validate_board.py
│     └─ validate.py
└─ tests/
   ├─ test_smoke.py
   ├─ test_cli.py
   ├─ test_documentation_board.py
   ├─ test_next_step.py
   ├─ test_handoff.py
   └─ test_validate.py
        """,
    )
    add_bullets(
        document,
        [
            "第一次读代码，先看 cli.py，了解有哪些命令。",
            "然后看 documentation_board.py、next_board.py、handoff_board.py、validate_board.py，理解每个命令如何编排。",
            "再看 parser.py 和 models.py，理解 Markdown 如何变成结构化对象。",
            "最后看 next_step.py、handoff.py、validate.py，理解真正的规则内容。",
        ],
    )

    document.add_heading("8. 如何把它用到你自己的项目里", level=1)
    add_paragraph(document, "这一步的关键是先把你的项目语义写进五个 Markdown 文件，而不是先改代码。")
    add_image(document, diagrams[4], "图 5  把它用到你自己的项目时怎么落地")
    add_numbered(
        document,
        [
            "进入你的项目目录后运行 cmb init。",
            "先写 AGENTS.md，明确协作边界和工作规则。",
            "再写 Plan.md，给出里程碑和可交付项。",
            "接着写 Documentation.md，把当前阶段、已完成事项、下一步、最近决策和最近验证填清楚。",
            "用 cmb status 和 cmb next 检查现在的状态是否清晰。",
            "开发过程中坚持写 cmb log，切线程时用 cmb handoff，阶段节点用 cmb validate。",
        ],
    )

    document.add_heading("9. 初学者最容易卡住的地方", level=1)
    add_bullets(
        document,
        [
            "把 Documentation.md 写成自由散文，结果 parser 无法稳定解析。",
            "把 Plan.md 里的任务写得过粗，导致 next 虽然能输出建议，但行动性不强。",
            "忘记维护 Latest Decision 和 Latest Verification，导致交接信息变得很弱。",
            "手工改坏标题层级，比如把 ### Current Phase 改成 ## Current Phase，结果 status 和 validate 同时受影响。",
        ],
    )
    add_paragraph(document, "遇到问题时，优先回到文件结构本身检查，通常比先看代码更快。")

    document.add_heading("10. 维护者和老手可以重点看什么", level=1)
    add_bullets(
        document,
        [
            "想扩展命令时，先定义 Markdown 契约，再补模型、解析、规则和 CLI。",
            "想保持命令稳定时，优先守住 store.py 只负责读写、parser.py 只负责解析的边界。",
            "想让 handoff 更稳时，先保证 Documentation.md 的内容质量，因为 handoff 不会凭空补全未记录的信息。",
            "想让 validate 更有价值时，优先增加结构性检查，不要让它变成自动修复器。",
        ],
    )

    document.add_heading("11. 推荐的最小填写模板", level=1)
    add_paragraph(document, "下面这个 Documentation.md 片段，是很适合初学者先照着填的一版。")
    add_code_block(
        document,
        """
# Documentation

## Current Status

### Current Phase
API Integration

### Completed Items
- 完成 CLI 骨架
- 完成基础测试

### Next Step
- 接入真实 API 并补充失败场景测试

### Latest Decision
- Decision: 保持规则驱动，不引入复杂推理
- Reason: 便于长期维护和排查

### Latest Verification
- Command: pytest -q
- Result: 20 passed

## Log Entries

### 2026-04-16 10:00:00
- Item: 完成状态读取命令
- Decision: 固定标题结构
- Reason: 方便解析和校验
- Verification Command: pytest -q
- Verification Result: 20 passed
        """,
    )

    document.add_heading("12. 最后的理解方式", level=1)
    add_paragraph(document, "如果你要用一句话记住这个项目，可以记成：它把项目当前状态写入固定文件，再用一组命令围绕这些文件进行读取、写回、建议、交接和自检。")
    add_paragraph(document, "对新手来说，它提供了清晰入口；对老手来说，它提供了稳定契约和明确边界。")
    add_paragraph(document, "只要 Plan.md 和 Documentation.md 写得清楚，这套工具就会很好用。", bold=True)

    return document


def main() -> None:
    ensure_output_dirs()
    diagrams = [
        build_workflow_diagram(),
        build_architecture_diagram(),
        build_file_contract_diagram(),
        build_next_rule_diagram(),
        build_adoption_diagram(),
    ]
    document = build_document(diagrams)
    document.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")
    for diagram in diagrams:
        print(f"Created: {diagram}")


if __name__ == "__main__":
    main()
